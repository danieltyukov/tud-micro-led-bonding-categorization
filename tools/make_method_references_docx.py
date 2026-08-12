#!/usr/bin/env python3
"""Build the method-reference note for the article as a .docx.

Same content as docs/article/METHOD_REFERENCES.md, in the format Ahmed works in.
IEEE-only reference set, every entry downloaded from Xplore and read.
Styles come from the existing section document at the repo root; the body is written fresh.
"""
import os
from docx import Document
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "article", "Method_references_and_draft_notes.docx")

doc = Document(os.path.join(ROOT, "Electrical_Characterization_section.docx"))
body = doc.element.body
for child in list(body):
    if not child.tag.endswith("}sectPr"):
        body.remove(child)


def P(text="", italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    if size:
        r.font.size = Pt(size)
    return p


def H(text, level):
    return doc.add_heading(text, level=level)


def REF(tag, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(34)
    p.paragraph_format.first_line_indent = Pt(-34)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(f"[{tag}] ").bold = True
    p.add_run(text)
    return p


def B(text):
    return doc.add_paragraph(text, style="List Bullet")


P("Method references and notes on the draft", bold=True, size=16)
P('For "Contactless Die Attach Methods Via Solder Paste", electrical characterization '
  "only. Prepared 12 August 2026.", italic=True)
P("Five references, IEEE only. Each was found on IEEE Xplore through the TU Delft Library "
  "proxy, the metadata was taken from the Xplore record, and the full text was downloaded "
  "and read before being tied to a sentence.")
P("Five rather than fourteen, because these are the ones that actually carry a claim we "
  "make. Where no IEEE source genuinely supports a method, that is said below rather than "
  "papered over with a citation that does not fit.")

H("1. Reference list", 1)
REFS = [
    ("R1", "A. Abdelwahab, H. van Zeijl, R. van Hoorn, H. Kuipers, and M. Mastrangeli, "
           "“Pick-and-Release: A Novel Contactless Bonding Method for Die Attachment,” "
           "in 2025 IEEE 75th Electronic Components and Technology Conference (ECTC), "
           "May 2025, pp. 2125-2132, doi: 10.1109/ECTC51687.2025.00363."),
    ("R2", "M. Zhanghu, Y. Liu, B.-R. Hyun, Y. Li, and Z. Liu, “Optimizing InGaN "
           "Micro-LED Efficiency: Investigating the Internal Quantum Efficiency and "
           "Ideality Factor Connection,” IEEE Trans. Electron Devices, vol. 71, no. 10, "
           "pp. 6190-6197, Oct. 2024, doi: 10.1109/TED.2024.3449829."),
    ("R3", "E. Jung, J. K. Lee, M. S. Kim, and H. Kim, “Leakage Current Analysis of "
           "GaN-Based Light-Emitting Diodes Using a Parasitic Diode Model,” IEEE Trans. "
           "Electron Devices, vol. 62, no. 10, pp. 3322-3325, Oct. 2015, "
           "doi: 10.1109/TED.2015.2468581."),
    ("R4", "N. Roccato et al., “Fast Characterization of Power LEDs: Circuit Design and "
           "Experimental Results,” IEEE Trans. Electron Devices, vol. 71, no. 6, "
           "pp. 3753-3760, Jun. 2024, doi: 10.1109/TED.2024.3393448."),
    ("R5", "D. Gacio, J. M. Alonso, J. Garcia, M. S. Perdigao, E. Sousa Saraiva, and "
           "F. E. Bisogno, “Effects of the Junction Temperature on the Dynamic "
           "Resistance of White LEDs,” IEEE Trans. Ind. Appl., vol. 49, no. 2, "
           "pp. 750-760, Mar./Apr. 2013, doi: 10.1109/TIA.2013.2243092."),
]
for tag, text in REFS:
    REF(tag, text)

H("2. Where each one goes, and what it actually says", 1)

H("[R1] Abdelwahab et al., ECTC 2025", 2)
P("Cite at the daisy-chain sentence and at the sheet-resistance sentence.")
P("This is the method source for both, not a background citation. Its Section E describes "
  "the same structure the new figure measures: a daisy chain bridging six 1 mm x 1 mm "
  "dummy dies on a PCB that also carries van der Pauw and TLM structures and RGB-LED "
  "footprints, read on a Summit 11K/12K probe station, with sheet resistance taken on the "
  "CDE ResMap 178 multi-probe station. It reports chain resistances from 0.199 ± 0.019 Ω "
  "(Ag, with pressure) to 0.380 ± 0.025 Ω (Au, pressure-less), the same order as the 0.22 "
  "to 0.61 Ω in the new eight-condition data.")
P("Because it uses the same instrument and the same structure, it removes the need to cite "
  "a generic four-point-probe or daisy-chain standard for either sentence.")

H("[R2] Zhanghu et al., TED 2024", 2)
P("Cite where the ideality window of 1.2 to 2.4 is defined, and where the anomalous "
  "channel is excluded.")
P("It supplies the physical reading of the number: n = 1 corresponds to band-to-band "
  "radiative recombination, n = 2 to Shockley-Read-Hall recombination through defect "
  "levels, and n > 2 to deep-level-assisted tunnelling. That is exactly the reasoning "
  "behind treating a fit outside the window as evidence of an extra conduction path rather "
  "than a bad fit. It also reports an inverse correlation between ideality factor and "
  "internal quantum efficiency, and observes that leakage paths raise the apparent series "
  "resistance, which is the same coupling we see on S3-D1.")

H("[R3] Jung et al., TED 2015", 2)
P("Cite where the parallel path across the junction is introduced, in the forward-sweep "
  "defect section and again in the reverse-bias section.")
P("It analyses LED leakage as a parasitic element shunting the main diode and separates "
  "the two by their turn-on voltages, 2.64 V for the main diode against 0.94 V for the "
  "parasitic one. That is the same separation we make when a channel conducts at 1.21 V, "
  "well below a red LED’s turn-on.")
P("One caveat to respect in the wording: their shunt is intrinsic to the device, "
  "attributed to hydrogen-related deep levels, while ours is external contamination from "
  "the assembly. Cite it for the model and the diagnostic, not for the mechanism.")

H("[R4] Roccato et al., TED 2024", 2)
P("Cite at the 5 ms current-on time and in the self-heating section.")
P("Its abstract states the point outright: LED characterization is often carried out with "
  "pulses of 10 ms and longer, conditions in which self-heating can significantly affect "
  "the measurement. Our own result, that nothing shifts between 5 and 20 ms but 80 ms "
  "moves the extracted series resistance, sits exactly on that boundary, so this is "
  "corroboration rather than background.")

H("[R5] Gacio et al., TIA 2013", 2)
P("Cite at the forward-voltage temperature coefficient and in the self-heating analysis.")
P("It defines the relation we use, V_D = V_D(25 °C) + λ(T_j − 25) with λ = ∂V_D/∂T_j, and "
  "measures λ experimentally for four commercial LEDs. More usefully for us, it then "
  "measures how junction temperature shifts the LED’s dynamic resistance, which is the "
  "exact confound behind our 80 ms result. It also picks its own pulse width as a "
  "trade-off between signal-to-noise and negligible self-heating, the same argument we "
  "make for 5 ms.")

H("3. What has no IEEE source, and what to do about it", 1)
P("Two things in the section cannot honestly be cited to IEEE.")
P("The statistics. The Wilson score interval on the yield bars, the chi-square test of "
  "independence across conditions, and the one-way ANOVA on series resistance are general "
  "statistics with no IEEE home. Three options, in order of preference:")
B("State them without citation. They are standard and named explicitly, which is normal "
  "practice for a packaging paper.")
B("Cite the primary sources, which are not IEEE: E. B. Wilson, J. Amer. Statist. Assoc., "
  "vol. 22, no. 158, pp. 209-212, 1927 for the interval, and a design-of-experiments text "
  "for the ANOVA and chi-square.")
B("Drop the Wilson interval from the figure and show plain counts.")
P("I would take the first option for the tests and keep the Wilson interval named in the "
  "caption as it already is.")
P("The diode-equation fit itself. Fitting V = V₀ + n·V_T·ln(I) + I·R_s by nonlinear least "
  "squares is textbook. [R2] extracts ideality factors from measured I-V curves and [R5] "
  "fits forward-voltage data, so between them the practice is covered. Adding a separate "
  "extraction paper would be padding.")

H("4. One discrepancy the reading turned up", 1)
P("The draft says the effective resistivity of the Au surface finish was 10 × 10⁻⁸ Ω·m, "
  "measured on the CDE ResMap 178.")
P("[R1], from the same group and the same instrument, reports 2.86 × 10⁻⁸ Ω·m for Au, with "
  "3.59 × 10⁻⁸ for Cu and 2.74 × 10⁻⁸ for Ag. Bulk gold is 2.44 × 10⁻⁸ Ω·m, so 2.86 is a "
  "sensible thin-film value and 10 is four times bulk.")
P("Either the new wafer genuinely differs from the one in the ECTC paper, in which case "
  "the factor of 3.5 is worth a sentence, or the number has been mistyped. Worth checking "
  "against the raw ResMap output before submission, since [R1] is cited two lines earlier "
  "and a reader can put the two numbers side by side.")

H("5. Three numbers in the draft worth fixing", 1)
P("Checked against the raw data. The conclusions do not change; the arithmetic does.")

H("5.1 The shunt fraction: neither one third nor 61 %", 2)
P("Paragraph beginning “One numerical correction:”. Two things.")
P("First, that paragraph is an editorial note written in the first person and it is "
  "sitting in the body of the article. It needs to come out.")
P("Second, the number. The correct figure is about 44 %, and the reason both earlier "
  "figures missed it is the test current. This meter is not a 1 mA source. Its diode-mode "
  "test current was measured directly during the round 1 setup as 0.142 V across a 100 Ω "
  "0.1 % reference, so 1.42 mA, cross-checked against a 1.49 mA short-circuit figure.")
P("At the 1.781 V reading a 2.87 kΩ shunt carries 0.62 mA. Against 1.42 mA that is 44 %, "
  "and the junction takes the remaining 0.80 mA.")
P("That split is confirmed independently. A healthy die on the same coupon sits at 1.749 V "
  "at 0.349 mA; carrying it to 0.80 mA through its own fitted ideality gives 1.783 to "
  "1.793 V, against the 1.781 V actually read. The alternative reading, that the meter is "
  "a Thevenin source of 3.245 V behind about 2.2 kΩ, would put only 0.05 mA through the "
  "junction and predict a reading roughly 150 mV low, so the data rules it out.")
P("Suggested replacement: “At the 1.781 V reading, the shunt carries 0.62 mA of the "
  "meter’s 1.42 mA diode-test current, so the junction sees only 0.80 mA and the channel "
  "still reads as a normal red LED.”", italic=True)

H("5.2 The variance share is 85 %, not 83 %", 2)
P("The re-seating paragraph currently reads “approximately 83 %, or about 85 %”. One "
  "number, and it is 85 %.")
P("The pooled within-condition standard deviation over the 29 physical red-channel fits is "
  "0.9147 Ω and the re-seating pooled standard deviation is 0.843 Ω, so the variance ratio "
  "is 0.843² / 0.9147² = 84.9 %. The 83 % comes from rounding 0.84 before squaring.")

H("5.3 The self-heating estimate does not reproduce 1.2 Ω", 2)
P("The stated inputs are 200 K/W, 28.4 mW and −2 mV/K. Because dissipated power is very "
  "nearly proportional to current over this range, the thermal perturbation is linear in "
  "current and folds entirely into the fitted series resistance:")
P("ΔR_s ≈ λ · R_th · V_F = (2 mV/K)(200 K/W)(2.01 V) = 0.80 Ω", italic=True)
P("not 1.2 Ω. Against the measured 1.49 Ω that is a factor of 1.9, so agreement “to within "
  "25 %” does not hold as written.")
P("Cleanest fix is to invert it: the measured 1.49 Ω implies R_th ≈ 370 K/W, ordinary for "
  "a 0404 package on two-layer FR-4 with no heat-spreading copper. [R5] supports doing it "
  "this way round, since it measures the temperature dependence of LED dynamic resistance "
  "directly. The conclusion is unaffected: heating sets in between 20 and 80 ms, so 5 ms "
  "is safe.")

H("6. Smaller editorial points", 1)
for txt in [
    "The equation symbols dropped out of the paraphrase in two places: “where  is the "
    "terminal voltage,  is the voltage offset,  is the ideality factor” and “Over this "
    "limited range , , and  were strongly correlated”.",
    "“a statistically significant association between assembly condition and channel "
    "p = 2.2 × 10⁻⁴” is missing a word after “channel”.",
    "Two cross-references still point at Section D.",
    "The sheet-resistance sentence gives resistivity but not the sheet resistance in Ω/sq "
    "or the Au thickness. Since ρ = R_sheet · t, giving two of the three lets a reader "
    "check it. See section 4 above.",
]:
    B(txt)

H("Note on [R1]", 1)
P("The title we had in our own bibliography was wrong. The Xplore record gives "
  "“Pick-and-Release: A Novel Contactless Bonding Method for Die Attachment”, not "
  "“Electrical Characterization of Capillary Self-Assembled Micro-LEDs on PCB "
  "Substrates”. The DOI was right. Worth checking anywhere else it is cited.")

doc.save(OUT)
print("wrote", OUT)
