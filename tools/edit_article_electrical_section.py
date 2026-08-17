#!/usr/bin/env python3
"""Apply Daniel's edits to the electrical-characterization section of Ahmed's article.

Input :  docs/article/round2/Article_from_Ahmed_2026-08-17.docx
Output:  docs/article/round2/Article_edited_DanielTyukov_2026-08-17.docx

Scope is the Electrical Characterization section only. Nothing outside it is touched.

Note on two things I earlier believed were broken and are not. The inline symbols in the
equation legend are OMML objects, invisible to text extraction but rendering correctly in
Word, so they are left alone. The Section B/C/D cross-references are also correct: the
subsections letter as A Channel screening, B Forward characteristics, C Defect detection,
D Measurement repeatability, E Self-heating, F Reverse bias, G Summary. Both were checked
by rendering the document and reading the result rather than the extracted text.

Every edit is recorded in CHANGES.md next to the output.
"""
import copy
import os

import docx
from docx.oxml.ns import qn
from docx.shared import RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "docs", "article", "round2",
                   "Article_from_Ahmed_2026-08-17.docx")
DST = os.path.join(ROOT, "docs", "article", "round2",
                   "Article_edited_DanielTyukov_2026-08-17.docx")
CHANGES = os.path.join(ROOT, "docs", "article", "round2", "CHANGES.md")

RED = RGBColor(0xFF, 0x00, 0x00)
MATH = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"

doc = docx.Document(SRC)
P = doc.paragraphs
log = []


def guard_no_math(par, where):
    """Rebuilding a paragraph that carries an OMML object would orphan the object."""
    if par._element.findall(".//" + MATH):
        raise SystemExit(f"REFUSING to rebuild {where}: it contains inline OMML math")


def build(par, segments, note):
    """Replace a paragraph's content with `segments`, keeping its style."""
    guard_no_math(par, note.split()[0])
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    for text, kind in segments:
        run = par.add_run(text)
        k = kind.split("+") if kind else []
        if "i" in k:
            run.italic = True
        if "sub" in k:
            run.font.subscript = True
        if "red" in k:
            run.font.color.rgb = RED
    log.append(note)


def sub_in_run(par, old, new, note):
    """Replace text inside whichever run holds it, leaving fields and math intact."""
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            log.append(note)
            return
    raise SystemExit(f"NOT FOUND: {old!r} in {par.text[:70]!r}")


def cite(par, anchor, tag, note):
    """Insert a red [Rn] immediately after `anchor`."""
    for r in par.runs:
        if anchor in r.text:
            head, _, tail = r.text.partition(anchor)
            r.text = head + anchor
            new = par.add_run(" " + tag)
            new.font.color.rgb = RED
            par._p.insert(list(par._p).index(r._element) + 1, new._element)
            if tail:
                rest = copy.deepcopy(r._element)
                for t in rest.findall(qn("w:t")):
                    t.text = tail
                    t.set(qn("xml:space"), "preserve")
                par._p.insert(list(par._p).index(new._element) + 1, rest)
            log.append(note)
            return
    raise SystemExit(f"ANCHOR NOT FOUND: {anchor!r} in {par.text[:70]!r}")


# --------------------------------------------------------------------------
# Daisy chain and sheet resistance
# --------------------------------------------------------------------------
build(P[31], [
    ("A simple PCB was designed and fabricated to verify the electrical continuity of "
     "dies attached under the different conditions. As shown in Fig. XX, the DC test "
     "structure consisted of a daisy chain incorporating six 1 × 1 mm² Au-coated dummy "
     "dies", ""),
    (" [R1]", "red"),
    (". The sheet resistance of the 4-inch Au-coated wafer was measured using a CDE "
     "ResMap 178 multiprobe station", ""),
    (" [R1]", "red"),
    (", and the corresponding effective resistivity of the Au surface finish was "
     "2.86 × 10⁻⁸ Ω·m. The total daisy-chain resistance measured for each assembly "
     "condition is summarized in Fig. XX.", ""),
], "31  daisy-chain and sheet-resistance sentences: added [R1] twice, closed the gap in "
   "'2.86× 10⁻⁸'")

sub_in_run(P[36], "Total chain resistance for all assembly conditions",
           "Total daisy-chain resistance for each assembly condition, measured across a "
           "chain of six 1 × 1 mm² Au-coated dummy dies. Markers are the mean and bars "
           "the deviation across the measured chains.",
           "36  daisy-chain figure caption expanded to say what was measured")

# --------------------------------------------------------------------------
# Chi-square sentence, missing word
# --------------------------------------------------------------------------
build(P[42], [
    ("The resulting electrical yield clearly differentiated the assembly conditions. A "
     "chi-square test applied to all 120 channels confirmed a statistically significant "
     "association between assembly condition and channel outcome (p = 2.2 × 10⁻⁴).", ""),
], "42  'between assembly condition and channel p = ...' was missing a word. Now 'channel "
   "outcome (p = ...)'")

# --------------------------------------------------------------------------
# Pulsed sweep
# --------------------------------------------------------------------------
build(P[60], [
    ("Each functional channel was characterized at 63 current levels ranging from 0.5 to "
     "15.9 mA. The current was supplied through six binary-weighted resistors connected to "
     "digital output pins, providing 2⁶ − 1 = 63 non-zero current combinations. To "
     "minimize self-heating and ensure a consistent thermal history across all measurement "
     "points, the current was pulsed for 5 ms followed by an off-time of 250 ms", ""),
    (" [R4]", "red"),
    (". The series resistance was extracted using a three-parameter nonlinear "
     "least-squares fit applied to all data points above 0.5 mA:", ""),
], "60  5 ms pulsing: added [R4], which reports self-heating bias for pulses of 10 ms and "
   "longer. Also spaced '(2⁶−1 =63)' as '2⁶ − 1 = 63'")

# Equation legend: the symbols are OMML and render correctly, so only the stray full stop
# after 'series resistance' is touched.
sub_in_run(P[64], " is the series resistance.",
           " is the series resistance,",
           "64  equation legend: 'series resistance. which is the sum' had a full stop "
           "mid-sentence, changed to a comma. The symbols were left untouched")

# --------------------------------------------------------------------------
# Defect detection. The shunt fraction was wrong in both earlier versions.
# --------------------------------------------------------------------------
build(P[81], [
    ("One red channel (S3–D1) produced a non-physical fit, with an ideality factor of "
     "13.0 and a negative series resistance. Below 3 mA, the channel conducted 0.42 mA at "
     "approximately 1.21 V, which is below the typical turn-on voltage of a red LED. This "
     "behavior indicates the presence of a parallel leakage path with an estimated "
     "resistance of approximately 2.9 kΩ, allowing current to bypass the LED junction", ""),
    (" [R3]", "red"),
    (". Above 3 mA, however, its current–voltage characteristic became nearly "
     "indistinguishable from that of a functional channel. The channel also passed the "
     "initial DMM screening, showing a forward voltage of 1.781 V at the meter's 1.42 mA "
     "diode-test current. At that reading the shunt carries approximately 0.62 mA, so only "
     "0.80 mA reaches the junction, yet the channel still registers as a normal red LED. A "
     "conventional diode test may therefore fail to detect this type of parallel leakage.",
     ""),
], "81  shunt fraction corrected. The meter sources 1.42 mA in diode mode, not 1 mA, so "
   "the shunt takes 0.62 mA and the junction 0.80 mA. Added [R3]")

cite(P[82], "ideality factors between 1.2 and 2.4", "[R2]",
     "82  ideality window: added [R2], which reads n = 1 as radiative, n = 2 as SRH "
     "through defect levels and n > 2 as deep-level-assisted tunnelling")

note83 = P[83]
log.append("83  DELETED the paragraph beginning 'One numerical correction:'. It was an "
           "editing note written in the first person, and its 61 % figure assumed a 1 mA "
           "test current")
note83._element.getparent().remove(note83._element)

# --------------------------------------------------------------------------
# Repeatability. Indices shift by one after the deletion above.
# --------------------------------------------------------------------------
P = doc.paragraphs
build(P[88], [
    ("To quantify this contribution, three channels were measured three times each. "
     "Between successive measurements, the same two jumper wires were disconnected and "
     "re-seated, while all other experimental parameters remained unchanged. The resulting "
     "pooled standard deviation was 0.84 Ω, compared with a within-condition standard "
     "deviation of 0.92 Ω. Based on the corresponding variances, jumper re-seating "
     "accounted for approximately 85 % of the variation that might otherwise have been "
     "interpreted as die-to-die resistance variation.", ""),
], "88  variance share: the draft gave both 83 % and 85 %. From the unrounded standard "
   "deviations, 0.843² / 0.9147² = 84.9 %, so 85 %")

# --------------------------------------------------------------------------
# Self-heating
# --------------------------------------------------------------------------
cite(P[95], "rather than assumed to be negligible", "[R4]",
     "95  pulse-duration argument: added [R4]")

build(P[99], [
    ("The extracted series resistance remained unchanged within the fitting uncertainty "
     "when the current-on time was increased from 5 to 20 ms. At 80 ms, however, it "
     "decreased by 1.49 Ω. Because the dissipated power is nearly proportional to current "
     "over this range, the thermal perturbation is itself nearly linear in current and is "
     "therefore absorbed almost entirely into the fitted series resistance, whose apparent "
     "shift is the product of the forward-voltage temperature coefficient, the "
     "junction-to-board thermal resistance and the forward voltage. Taking a coefficient "
     "of approximately −2 mV/K, the measured shift implies a junction-to-board thermal "
     "resistance of approximately 370 K/W, corresponding to a junction-temperature rise of "
     "about 11 K at 28.4 mW. This is a plausible value for a 0404 package on two-layer "
     "FR-4 with no heat-spreading copper", ""),
    (" [R5]", "red"),
    (".", ""),
], "99  self-heating estimate rewritten. The stated inputs (200 K/W, −2 mV/K, 2.01 V) give "
   "0.80 Ω, not 1.2 Ω, so the claimed agreement to within 25 % did not hold. Inverted the "
   "calculation instead: the measured 1.49 Ω implies about 370 K/W. Added [R5]")

# --------------------------------------------------------------------------
# Reverse bias
# --------------------------------------------------------------------------
cite(P[102], "conduction paths parallel to the LED junction", "[R3]",
     "102 reverse-bias check: added [R3]")

sub_in_run(P[104], "initial 1 mA diode test", "initial 1.42 mA diode test",
           "104 diode-test current corrected to 1.42 mA here as well")

# --------------------------------------------------------------------------
# Measurement-setup table
# --------------------------------------------------------------------------
for row in doc.tables[3].rows:
    cell = row.cells[3]
    if "≈1 mA forward" in cell.text:
        for par in cell.paragraphs:
            for r in par.runs:
                if "≈1 mA forward" in r.text:
                    r.text = r.text.replace("≈1 mA forward", "1.42 mA forward")
                    log.append("T3  measurement-setup table: DMM diode-test excitation "
                               "corrected from ≈1 mA to 1.42 mA, measured as 0.142 V "
                               "across a 100 Ω 0.1 % reference")

# --------------------------------------------------------------------------
# Reference key, appended in red so it can be assigned in EndNote and deleted.
# --------------------------------------------------------------------------
REFS = [
    ("R1", "A. Abdelwahab, H. van Zeijl, R. van Hoorn, H. Kuipers, and M. Mastrangeli, "
           "“Pick-and-Release: A Novel Contactless Bonding Method for Die "
           "Attachment,” in 2025 IEEE 75th Electronic Components and Technology "
           "Conference (ECTC), May 2025, pp. 2125-2132, doi: 10.1109/ECTC51687.2025.00363."),
    ("R2", "M. Zhanghu, Y. Liu, B.-R. Hyun, Y. Li, and Z. Liu, “Optimizing InGaN "
           "Micro-LED Efficiency: Investigating the Internal Quantum Efficiency and "
           "Ideality Factor Connection,” IEEE Trans. Electron Devices, vol. 71, "
           "no. 10, pp. 6190-6197, Oct. 2024, doi: 10.1109/TED.2024.3449829."),
    ("R3", "E. Jung, J. K. Lee, M. S. Kim, and H. Kim, “Leakage Current Analysis of "
           "GaN-Based Light-Emitting Diodes Using a Parasitic Diode Model,” IEEE "
           "Trans. Electron Devices, vol. 62, no. 10, pp. 3322-3325, Oct. 2015, "
           "doi: 10.1109/TED.2015.2468581."),
    ("R4", "N. Roccato et al., “Fast Characterization of Power LEDs: Circuit Design "
           "and Experimental Results,” IEEE Trans. Electron Devices, vol. 71, no. 6, "
           "pp. 3753-3760, Jun. 2024, doi: 10.1109/TED.2024.3393448."),
    ("R5", "D. Gacio, J. M. Alonso, J. Garcia, M. S. Perdigao, E. Sousa Saraiva, and "
           "F. E. Bisogno, “Effects of the Junction Temperature on the Dynamic "
           "Resistance of White LEDs,” IEEE Trans. Ind. Appl., vol. 49, no. 2, "
           "pp. 750-760, Mar./Apr. 2013, doi: 10.1109/TIA.2013.2243092."),
]
head = doc.add_paragraph()
hr = head.add_run("Reference key for [R1] to [R5], for EndNote. Delete this block once the "
                  "entries are assigned.")
hr.bold = True
hr.font.color.rgb = RED
for tag, text in REFS:
    par = doc.add_paragraph()
    run = par.add_run(f"[{tag}] {text}")
    run.font.color.rgb = RED
log.append("END appended a red reference key listing [R1] to [R5] in full, to be deleted "
           "once the entries are in EndNote")

doc.save(DST)

with open(CHANGES, "w") as fh:
    fh.write("# Edits to the electrical-characterization section\n\n")
    fh.write("Daniel Tyukov, 17 August 2026. Source: `Article_from_Ahmed_2026-08-17.docx`. "
             "Output: `Article_edited_DanielTyukov_2026-08-17.docx`.\n\n")
    fh.write("Numbers are paragraph indices in the source file. Citations are red, "
             "everything else is plain so it reads as normal body text.\n\n")
    for line in log:
        fh.write(f"- {line}\n")
    fh.write("\n## Checked and deliberately left alone\n\n")
    fh.write("- The inline symbols in the equation legend, and in the sentence about the "
             "three correlated parameters. They are OMML objects: invisible to text "
             "extraction, correct in Word. I flagged these as missing in my previous mail. "
             "They are not missing.\n")
    fh.write("- The Section B, C and D cross-references. They are correct as written. The "
             "subsections letter A Channel screening, B Forward characteristics, C Defect "
             "detection, D Measurement repeatability, E Self-heating, F Reverse bias, "
             "G Summary. I flagged these too, also wrongly.\n")

print("wrote", DST)
print("wrote", CHANGES)
print(f"{len(log)} edits")
for line in log:
    print("  -", line)
