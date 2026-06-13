#!/usr/bin/env python3
"""
Generate 'Electrical characterization' section as a Word .docx for Ahmed's paper.

- Filled reference tables (instruments / structures).
- Empty measurement tables with bordered, fillable cells.
- Unicode / run-level subscripts and superscripts (no LaTeX, no equation editor).

Usage:  python3 tools/make_electrical_char_docx.py
Output: Electrical_Characterization_section.docx  (repo root)
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Times New Roman"
BODY_SIZE = 11
TABLE_SIZE = 9.5

# ----------------------------------------------------------------------------
# Rich-text mini-markup:  R_{sh}  -> R + subscript "sh"
#                          cm^{2}  -> cm + superscript "2"
# Everything else (Unicode α ρ Ω π × · µ ≈ → …) passes through literally.
# ----------------------------------------------------------------------------
_TOKEN = re.compile(r"(_\{[^}]*\}|\^\{[^}]*\})")


def add_rich(paragraph, text, *, bold=False, italic=False, size=None, color=None):
    """Append run(s) to paragraph, honouring _{..} / ^{..} sub/superscript markup."""
    for tok in _TOKEN.split(text):
        if not tok:
            continue
        sub = sup = False
        body = tok
        if tok.startswith("_{") and tok.endswith("}"):
            body, sub = tok[2:-1], True
        elif tok.startswith("^{") and tok.endswith("}"):
            body, sup = tok[2:-1], True
        run = paragraph.add_run(body)
        run.font.name = BODY_FONT
        run.font.size = Pt(size if size else BODY_SIZE)
        run.bold = bold
        run.italic = italic
        run.font.subscript = sub
        run.font.superscript = sup
        if color is not None:
            run.font.color.rgb = color
    return paragraph


def para(doc, text="", *, bold=False, italic=False, size=None, color=None,
         space_after=6, align=None, first_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if text:
        add_rich(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def equation(doc, text):
    """Centred display equation in its own paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    add_rich(p, text, italic=False)
    return p


def heading(doc, text, level):
    h = doc.add_heading("", level=level)
    add_rich(h, text)
    for run in h.runs:
        run.font.name = BODY_FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# ----------------------------------------------------------------------------
# Table helpers
# ----------------------------------------------------------------------------
def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _fill_cell(cell, text, *, bold=False, header=False, align="left", size=TABLE_SIZE):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        add_rich(p, text, bold=bold or header, size=size)
    if header:
        _set_cell_bg(cell, "D9E2F3")  # light blue header


def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    add_rich(p, text, bold=True, size=9.5)
    return p


def make_table(doc, headers, rows, *, col_align=None, empty_rows=0,
               first_col_align="left"):
    """headers: list[str]; rows: list[list[str]] (use '' for blank fillable cells)."""
    ncol = len(headers)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # header
    for j, htext in enumerate(headers):
        _fill_cell(table.rows[0].cells[j], htext, header=True, align="center")
    # data rows
    all_rows = list(rows) + [[""] * ncol for _ in range(empty_rows)]
    for r in all_rows:
        cells = table.add_row().cells
        for j in range(ncol):
            val = r[j] if j < len(r) else ""
            a = first_col_align if j == 0 else (col_align[j] if col_align else "center")
            _fill_cell(cells[j], val, align=a)
    return table


# ============================================================================
# BUILD DOCUMENT
# ============================================================================
doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(BODY_SIZE)

for lvl, sz in ((1, 15), (2, 12.5), (3, 11.5)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = BODY_FONT
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0, 0, 0)

# ---- Editorial note (delete before submission) -----------------------------
note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(10)
add_rich(note,
         "[ Working note — delete before submission. ] Tables II–VIII are "
         "pre-formatted with empty cells; fill them once the cleanroom "
         "measurements are in. The first column of every measurement table is "
         "labelled “Sample / condition” so the same section serves "
         "both papers: for the molten-solder paper put the paste alloy and pad "
         "finish there (e.g. SAC305 / ENIG, TS391LT Sn42Bi57.6Ag0.4 / ENIG); "
         "for the liquid-assisted-assembly paper put the assembly medium or "
         "process variant. Symbols are real Word sub/superscripts, so they "
         "edit like normal text. Citation placeholders read “[ref]”.",
         italic=True, size=9.5, color=RGBColor(0x8B, 0x00, 0x00))

# ---- Section title ---------------------------------------------------------
heading(doc, "Electrical characterization", 1)

para(doc,
     "Once the micro-LEDs are bonded, each populated coupon is characterized "
     "electrically to confirm that the joints carry current and to quantify "
     "joint quality independently of the optical result. Five quantities are "
     "extracted on the same ENIG substrate and, where possible, on the same "
     "coupon: daisy-chain bond resistance and yield, the LED forward voltage "
     "and current–voltage (I–V) characteristic, the specific contact "
     "resistivity from transmission-line structures, the sheet resistance from "
     "van der Pauw cloverleaves, and the junction-to-substrate thermal "
     "resistance from forward-voltage thermometry. Impedance spectroscopy is "
     "an optional sixth measurement that resolves the joint into an equivalent "
     "circuit. All structures are contacted with four-wire (Kelvin) probes to "
     "remove lead and contact resistance from the reading. Table I lists the "
     "instruments and the structure used for each quantity.")

caption(doc, "Table I. Measurement setup: quantity, on-board test structure, "
             "instrument, and excitation.")
make_table(
    doc,
    ["Quantity", "Test structure", "Instrument", "Excitation / conditions"],
    [
        ["Bond / chain resistance",
         "Daisy chains DCL6, DCL12 (2N bonds)",
         "Keithley 2450 SourceMeter, 4-wire",
         "DC, I-sweep"],
        ["Forward voltage & I–V",
         "Bonded LED, four bottom contacts",
         "Keithley 2450 SourceMeter",
         "I = 10 mA (sense); ± sweep"],
        ["Specific contact resistivity ρ_{c}",
         "TLM ladders, W ∈ {0.25, 0.5, 1.0} mm",
         "HP 4156 parameter analyzer",
         "4-point, gaps 0.2–4 mm"],
        ["Sheet resistance R_{sh}",
         "van der Pauw cloverleaves, W ∈ {1.0, 0.5, 0.25, 0.15} mm",
         "HP 4156 parameter analyzer",
         "4-point, 90° rotation"],
        ["Junction–substrate R_{th}",
         "LED + NTC TH1–TH4 (interleaved)",
         "Keithley 2450, pulsed sense",
         "100 µA, 1 ms, 10% duty"],
        ["Series / parallel impedance",
         "LED + EIS OPEN / SHORT / LOAD",
         "Solartron 1260A + 1296A",
         "10 mV AC, 0.1 Hz–1 MHz"],
    ],
    col_align=["left", "left", "left", "left"],
)

# ---- A. Daisy chain --------------------------------------------------------
heading(doc, "A.  Daisy-chain bond resistance and yield", 2)
para(doc,
     "The daisy chains DCL6 and DCL12 place 6 and 12 LEDs in series through "
     "the red cathode, so a single I–V sweep across the chain endpoints "
     "exercises 2N bonds (N LEDs, two bonds each). The mean resistance per "
     "bond follows from the measured chain resistance:")
equation(doc, "R_{bond} = R_{chain} / (2N)")
para(doc,
     "A defective bond appears either as a discontinuity in the chain I–V "
     "curve or as a chain forward voltage that overshoots N × V_{F,nom}. "
     "Yield is reported as the number of conducting bonds over the total. For "
     "reference, the v1 board reported chain resistances of 0.199–0.380 Ω "
     "across die finishes and mounting conditions [ref]; values of that order "
     "indicate well-formed joints.")
caption(doc, "Table II. Daisy-chain bond resistance and yield.")
make_table(
    doc,
    ["Sample / condition", "Chain", "N (LEDs)", "R_{chain} (Ω)",
     "R_{bond} (mΩ)", "Bonds OK / total", "Yield (%)"],
    [
        ["", "DCL6", "6", "", "", "", ""],
        ["", "DCL12", "12", "", "", "", ""],
        ["", "DCL6", "6", "", "", "", ""],
        ["", "DCL12", "12", "", "", "", ""],
    ],
    first_col_align="left",
)

# ---- B. Vf / IV ------------------------------------------------------------
heading(doc, "B.  Forward voltage and current–voltage characteristics", 2)
para(doc,
     "Each addressable LED is read across its four bonded contacts. The "
     "forward voltage at the 10 mA operating point, V_{F}(10 mA), is the "
     "primary number; for WL-SFCC red the nominal value is ≈ 2.05 V [ref], "
     "so a joint that adds appreciable series resistance shows up as an "
     "elevated V_{F}. A least-squares fit of the high-current branch gives the "
     "joint series resistance R_{s}, and the reverse branch at −5 V gives "
     "the leakage current I_{leak}, which flags damaged or contaminated dies.")
caption(doc, "Table III. Per-LED forward voltage, series resistance, and "
             "leakage. Colours: R = red, G = green, B = blue.")
make_table(
    doc,
    ["Sample / condition", "LED ID", "Colour", "V_{F} at 10 mA (V)",
     "R_{s} from I–V (Ω)", "I_{leak} at −5 V (µA)", "Notes"],
    [["", "", "", "", "", "", ""] for _ in range(5)],
    first_col_align="left",
)

# ---- C. TLM ----------------------------------------------------------------
heading(doc, "C.  Specific contact resistivity (transmission-line method)", 2)
para(doc,
     "The TLM ladders give the sheet resistance R_{sh} and the specific "
     "contact resistivity ρ_{c} from the total resistance measured between "
     "successive finger pairs at known gaps d. For identical contacts of width "
     "W, length L, and transfer length L_{T},")
equation(doc, "R_{tot}(d) = (R_{sh} / W) · d + 2R_{c}")
equation(doc, "R_{c} = (R_{sh} · L_{T} / W) · coth(L / L_{T})")
para(doc,
     "A linear fit of R_{tot} versus d gives R_{sh} from the slope and 2R_{c} "
     "from the intercept; the specific contact resistivity then follows as "
     "ρ_{c} = R_{c} · W · L_{T} · tanh(L / L_{T}). The three "
     "finger widths in the ladder make the width dependence of L_{T} directly "
     "measurable.")
caption(doc, "Table IV. TLM extraction per finger width.")
make_table(
    doc,
    ["Sample / condition", "Width W (mm)", "R_{sh} (Ω/sq)",
     "2R_{c} (Ω)", "L_{T} (µm)", "ρ_{c} (Ω·cm^{2})"],
    [
        ["", "0.25", "", "", "", ""],
        ["", "0.50", "", "", "", ""],
        ["", "1.00", "", "", "", ""],
    ],
    first_col_align="left",
)

# ---- D. VDP ----------------------------------------------------------------
heading(doc, "D.  Sheet resistance (van der Pauw)", 2)
para(doc,
     "Each symmetric cloverleaf gives the sheet resistance directly, with no "
     "contact-geometry correction, from the implicit van der Pauw relation")
equation(doc, "exp(−π R_{A} / R_{sh}) + exp(−π R_{B} / R_{sh}) = 1")
para(doc,
     "where R_{A} and R_{B} are the transverse four-point resistances measured "
     "at 90° rotation. For R_{A} ≈ R_{B} ≡ R this collapses to "
     "R_{sh} = (π / ln 2) · R. The cloverleaves cross-check the TLM "
     "slope and surface any width-dependent edge effects. The v1 board "
     "reported an effective sheet resistivity of 2.94 × 10⁻⁵ Ω·m "
     "for the ENIG pads [ref], a useful sanity check on the new readings.")
caption(doc, "Table V. van der Pauw sheet resistance per cloverleaf.")
make_table(
    doc,
    ["Sample / condition", "Cloverleaf W (mm)", "R_{A} (Ω)",
     "R_{B} (Ω)", "R_{sh} (Ω/sq)", "ρ_{sheet} (Ω·m)"],
    [
        ["", "1.00", "", "", "", ""],
        ["", "0.50", "", "", "", ""],
        ["", "0.25", "", "", "", ""],
        ["", "0.15", "", "", "", ""],
    ],
    first_col_align="left",
)

# ---- E. Vf-TSP thermal -----------------------------------------------------
heading(doc, "E.  Junction-to-substrate thermal resistance (forward-voltage "
             "thermometry)", 2)
para(doc,
     "The LED forward voltage at a fixed low sense current falls almost "
     "linearly with junction temperature, so the diode is its own "
     "thermometer. With the temperature-sensitive parameter "
     "S ≡ −∂V_{F}/∂T_{J} calibrated once,")
equation(doc, "T_{J}(t) = T_{J,0} − (1/S) · [ V_{F}(t) − V_{F,0} ]")
para(doc,
     "The substrate temperature is read from the interleaved NTCs (TH1–TH4) "
     "through the Steinhart-B relation 1/T = 1/T_{0} + (1/B) · ln(R/R_{0}), "
     "with T_{0} = 298.15 K and R_{0} = 10 kΩ. To suppress self-heating "
     "bias the sense current is pulsed (100 µA, 1 ms, 10% duty). The "
     "bond-joint thermal resistance is the steady-state ratio of the "
     "junction-to-substrate temperature difference to the dissipated power:")
equation(doc, "R_{th,bond} = [ T_{J}(∞) − T_{NTC}(∞) ] / P_{op}")
caption(doc, "Table VI. Forward-voltage thermometry: calibrated slope and "
             "extracted junction-to-substrate thermal resistance.")
make_table(
    doc,
    ["Sample / condition", "LED ID", "S = −∂V_{F}/∂T_{J} (mV/K)",
     "P_{op} (mW)", "ΔT (K)", "R_{th,bond} (K/W)"],
    [["", "", "", "", "", ""] for _ in range(4)],
    first_col_align="left",
)

# ---- F. EIS ----------------------------------------------------------------
heading(doc, "F.  Impedance spectroscopy (optional)", 2)
para(doc,
     "Sweeping a 10 mV AC excitation from 0.1 Hz to 1 MHz and recording the "
     "complex impedance Z(jω) = Z′ + jZ″ separates the series bond "
     "resistance (the high-frequency intercept, Z′ → R_{s} as "
     "ω → ∞) from the diode’s parallel RC dynamics. For the "
     "standard LED equivalent circuit (R_{s} in series with R_{p} ∥ C_{j}),")
equation(doc, "Z(jω) = R_{s} + R_{p} / (1 + jω R_{p} C_{j})")
para(doc,
     "The on-board EIS OPEN, SHORT, and LOAD references (the 100 Ω 0.1% "
     "resistor) calibrate the fixture before the LED is measured.")
caption(doc, "Table VII. EIS equivalent-circuit parameters.")
make_table(
    doc,
    ["Sample / condition", "LED ID", "R_{s} (Ω)", "R_{p} (Ω)",
     "C_{j} (nF)", "f_{peak} (kHz)"],
    [["", "", "", "", "", ""] for _ in range(4)],
    first_col_align="left",
)

# ---- G. Summary ------------------------------------------------------------
heading(doc, "G.  Summary of electrical results", 2)
para(doc,
     "Table VIII collects the headline number from each method into a single "
     "per-condition comparison. Cross-condition variance is the figure of "
     "merit: it captures how wide each assembly process window actually is.")
caption(doc, "Table VIII. Summary of electrical characterization across "
             "assembly conditions (one row per condition).")
make_table(
    doc,
    ["Condition", "Yield (%)", "Mean R_{bond} (mΩ)", "V_{F} at 10 mA (V)",
     "ρ_{c} (Ω·cm^{2})", "R_{sh} (Ω/sq)", "R_{th,bond} (K/W)"],
    [["", "", "", "", "", "", ""] for _ in range(5)],
    first_col_align="left",
)

out = "/home/danieltyukov/workspace/tud/tud-micro-led-bonding-categorization/Electrical_Characterization_section.docx"
doc.save(out)
print("wrote", out)
