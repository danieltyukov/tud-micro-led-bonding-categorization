#!/usr/bin/env python3
"""Fill Electrical_Characterization_section.docx with round 1 and round 2 results."""
import copy, csv, math, os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FIG  = os.path.join(ROOT, "FINAL_MEASUREMENTS", "analysis", "figures")
OUT  = os.path.join(ROOT, "deliverable")
os.makedirs(OUT, exist_ok=True)

# ----------------------------------------------------------------- source data
fits = list(csv.DictReader(open(os.path.join(ROOT, "FINAL_MEASUREMENTS",
                                             "analysis", "fit_results.csv"))))
r1 = [r for r in csv.DictReader(open(os.path.join(ROOT, "RESULTS", "R1_channels.csv")))
      if r["die"].startswith("D") and not r["die"].startswith("DC")]

MODES = ["pass", "suspect", "cross_lit", "open", "short", "die_detached"]
N = {s: {m: 0 for m in MODES} for s in range(1, 9)}
for r in r1:
    N[int(r["sample"])][r["verdict"]] = N[int(r["sample"])].get(r["verdict"], 0) + 1
tot = {s: sum(N[s].values()) for s in N}
ok = {s: N[s]["pass"] for s in N}

red = [f for f in fits if f["colour"] == "R" and f["physical"] == "1"]
bys = {}
for f in red:
    bys.setdefault(int(f["sample"]), []).append(f)

def mean(v): return sum(v) / len(v)
def sd(v):
    return math.sqrt(sum((x - mean(v))**2 for x in v) / (len(v) - 1)) if len(v) > 1 else float("nan")

SEAT_SD = 0.843
DOMINANT = {1: "detached", 2: "detached / open", 3: "short / cross-lit", 4: "mixed",
            5: "none", 6: "detached", 7: "none", 8: "short"}

doc = Document(os.path.join(ROOT, "Electrical_Characterization_section.docx"))
SECT = {"C.  Specific contact resistivity": "D.  Specific contact resistivity",
        "D.  Sheet resistance": "E.  Sheet resistance",
        "E.  Junction-to-substrate": "F.  Junction-to-substrate",
        "F.  Impedance spectroscopy": "G.  Impedance spectroscopy",
        "G.  Summary of electrical": "H.  Summary of electrical"}

# ------------------------------------------------------------------- utilities
def para_by_text(frag):
    for p in doc.paragraphs:
        if frag.lower() in p.text.lower():
            return p
    raise KeyError(frag)

def insert_after(anchor, new_p):
    anchor._p.addnext(new_p._p)
    return new_p

def new_para(text="", style=None, italic=False, size=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if italic or size:
        for r in p.runs:
            if italic: r.italic = True
            if size:   r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p

def add_block(anchor, items):
    """items: list of (kind, payload). Returns the last inserted paragraph."""
    cur = anchor
    for kind, payload in items:
        if kind == "p":
            cur = insert_after(cur, new_para(payload))
        elif kind == "cap":
            cur = insert_after(cur, new_para(payload, italic=True, size=8.5,
                                            align=WD_ALIGN_PARAGRAPH.CENTER))
        elif kind == "fig":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(os.path.join(FIG, payload), width=Cm(8.5))
            cur = insert_after(cur, p)
    return cur

def fill(tbl, rows, keep_header=True):
    while len(tbl.rows) > (1 if keep_header else 0):
        tbl._tbl.remove(tbl.rows[-1]._tr)
    proto = None
    for r in rows:
        row = tbl.add_row()
        for c, val in enumerate(r):
            cell = row.cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
    return tbl

# --------------------------------------------------------- drop the working note
doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

# ------------------------------------------------------------------- section A
a = para_by_text("A defective bond appears either as a discontinuity")
add_block(a, [("p",
    "The chains could not be used. Both DCL6 and DCL12 read open in forward and reverse on "
    "every coupon, and gave no photovoltage under illumination. The cause is a layout fault "
    "rather than a bonding failure: the chain routes each die's top-left pad to its top-right "
    "pad, and with the dice rotated 90 degrees those two pads are the red and blue cathodes. "
    "Every element in the chain is therefore a pair of back-to-back diodes, and the real anode "
    "on the bottom-left corner connects to nothing. No chain resistance or chain yield can be "
    "extracted from this coupon revision.")])
fill(doc.tables[1], [["all", "DCL6", 6, "open", "n/a", "0 / 6", "n/a"],
                     ["all", "DCL12", 12, "open", "n/a", "0 / 12", "n/a"]])

# ------------------------------------------------------------------- section B
b = para_by_text("Each addressable LED is read across its four bonded contacts")
rows_iv = []
for s in sorted(bys):
    v = [float(f["Rs"]) for f in bys[s]]
    vf = [float(f["vf10"]) for f in bys[s]]
    rows_iv.append([f"{s}", f"{len(v)}", "R",
                    f"{mean(vf):.4f}" + (f" ± {sd(vf):.4f}" if len(v) > 1 else ""),
                    f"{mean(v):.2f}" + (f" ± {sd(v):.2f}" if len(v) > 1 else ""),
                    "not measured",
                    f"ideality {min(float(f['nid']) for f in bys[s]):.2f}–"
                    f"{max(float(f['nid']) for f in bys[s]):.2f}"])
t3 = doc.tables[2]
t3.rows[0].cells[1].text = "Channels n"
for r_ in t3.rows[0].cells[1].paragraphs[0].runs: r_.font.size = Pt(8)
fill(t3, rows_iv)

cur = add_block(b, [
 ("p", "Each channel was swept over 63 current levels between 0.5 and 15.9 mA using a "
       "binary-weighted resistor bank on six digital pins, with the current pulsed for 5 ms "
       "and off for 250 ms to hold the duty cycle fixed. Series resistance comes from a "
       "three-parameter least-squares fit of V = V0 + n·VT·ln(I) + I·Rs over all points above "
       "0.5 mA."),
 ("p", "Only the red channels give physical fits. Green and blue return ideality factors of "
       "3.2 to 4.0, which no LED has. With a forward voltage near 2.8 V the 5 V rail leaves "
       "only a 20-fold current range, over which V0, n and Rs trade off against each other and "
       "the fit becomes degenerate. Red spans 30-fold and returns ideality 1.60 to 2.06. All "
       "series resistance results below are red only."),
 ("fig", "fig1_iv_curves.png"),
 ("cap", "Fig. 1. Forward characteristics of one die (S1 D1) in all three colours. Red reaches "
         "13.9 mA; green and blue are limited to about 10 mA by the 5 V supply."),
 ("p", f"Across {len(red)} red channels the mean series resistance is {mean([float(f['Rs']) for f in red]):.2f} Ω "
       f"with a standard deviation of {sd([float(f['Rs']) for f in red]):.2f} Ω. One-way ANOVA "
       f"across the eight conditions gives F(7,21) = 1.73, p = 0.16. Forward voltage at 10 mA "
       f"behaves the same way: {mean([float(f['vf10']) for f in red]):.4f} V mean, "
       f"{sd([float(f['vf10']) for f in red])*1000:.1f} mV standard deviation, p = 0.24. "
       f"Neither quantity separates the assembly conditions."),
 ("fig", "fig4_rs_by_condition.png"),
 ("cap", "Fig. 2. Series resistance by condition. Red points are individual channels, squares "
         "are condition means, bars are ±1 standard deviation. The shaded band is the ±1σ "
         "measurement repeatability from Fig. 4."),
 ("p", "The limit is the fixture rather than the dice. Because this build takes both voltage "
       "taps from the breadboard instead of the Tier-1 probe pads, the two female-male jumpers "
       "and their header contacts sit inside the measured loop. Re-seating the same two jumpers "
       f"on an unchanged channel moves Rs by {SEAT_SD:.2f} Ω (pooled standard deviation, three "
       "channels, three seatings each), which accounts for 85 % of the variance otherwise read "
       "as die-to-die spread. Every re-seating also lands at or above the channel minimum, as "
       "contact resistance can only add, so the true red series resistance is near the mean of "
       "the minima, 10.6 Ω. Bond resistances are 10 to 100 mΩ. The smallest difference this "
       "arrangement can resolve at n = 4 is 1.4 Ω, so the null result above is a statement "
       "about the instrument and not about the bonds."),
 ("fig", "fig5_reseat_repeatability.png"),
 ("cap", "Fig. 3. Three channels, each measured after three independent re-seatings of the two "
         "jumpers. Nothing else changed between points."),
 ("p", "The sweep did find one defect the diode test missed. Channel S3 D1 red fits to an "
       "ideality of 13.0 and a negative series resistance, which is unphysical. Below 3 mA the "
       "junction sits at 1.21 V, under a red LED's turn-on, while 0.42 mA still flows: a "
       "parallel path of roughly 2.9 kΩ carries the current around the die. A diode test at "
       "1 mA reads 1.781 V and passes the channel. Any fit falling outside ideality 1.2 to 2.4 "
       "flags this condition, which makes the sweep a screen for shunts as well as a "
       "measurement."),
 ("fig", "fig7_shunt_detection.png"),
 ("cap", "Fig. 4. A contaminated die (S3 D1) against a healthy one. The curves agree above "
         "3 mA and diverge by 540 mV at the bottom of the sweep."),
])

# ----------------------------------------------- new section C: yield and modes
anchor_c = para_by_text("Table III. Per-LED forward voltage")
c_head = insert_after(anchor_c, new_para("C.  Channel yield and failure mode", style="Heading 2"))
rows_y = []
for s in range(1, 9):
    modes = ", ".join(f"{m.replace('die_detached','detached').replace('cross_lit','cross-lit')}={N[s][m]}"
                      for m in MODES[1:] if N[s][m]) or "none"
    rows_y.append([f"{s}", f"{ok[s]} / {tot[s]}", f"{100*ok[s]/tot[s]:.1f}", modes])

add_block(c_head, [
 ("p", "Every addressable channel on all five coupons was screened with a DMM diode test "
       "before any sweep: 120 channels, 8 conditions. A channel counts as passing if it shows "
       "a forward junction voltage and lights. The remaining verdicts record how it failed."),
 ("p", "This is the measurement that separates the conditions. A chi-square test on the 120 "
       "channels gives p = 2.2 × 10⁻⁴. Conditions 5 and 7 lost nothing at all. Condition 3 lost "
       "9 of 12 channels, and its failures are shorts and cross-lit pairs, which is what excess "
       "solder and pad-to-pad bridging look like. Conditions 1, 2 and 6 fail the other way, "
       "through detached dice and open cathodes, which is the signature of insufficient wetting "
       "or a weak bond. Condition 8 shows two shorts and nothing else."),
 ("fig", "fig2_defect_map.png"),
 ("cap", "Fig. 5. Channel-level screening result. Three rows per condition (red, green, blue), "
         "one column per die position. Blank cells are die positions the condition does not own."),
 ("fig", "fig3_yield_modes.png"),
 ("cap", "Fig. 6. Yield and failure mode by condition. Error bars are 95 % Wilson intervals on "
         "the pass fraction. Counts above each bar are passing channels over total."),
 ("p", "The failure mode carries more information than the yield number. Conditions splitting "
       "into an over-bonded group (3, 4, 8: shorts and bridges) and an under-bonded group "
       "(1, 2, 6: opens and detachment) puts conditions 5 and 7 inside the process window and "
       "the others on one side of it or the other. That ordering is consistent across 120 "
       "channels and does not depend on the fixture problem that limits the series-resistance "
       "comparison."),
])

# ------------------------------------------- sections not performed / partial
for frag, msg, tno in [
  ("A linear fit of Rtot versus d", "The TLM ladders were not measured. They require a "
   "four-wire source-measure unit on bare-ENIG structures, which was outside the instrument "
   "set used here.", 3),
  ("where RA and RB are the transverse", "The van der Pauw cloverleaves were not measured, "
   "for the same reason as the TLM ladders.", 4),
  ("The on-board EIS OPEN, SHORT, and LOAD", "Impedance spectroscopy was not performed. No LCR "
   "meter was available for this campaign.", 6)]:
    add_block(para_by_text(frag), [("p", msg)])
    fill(doc.tables[tno], [["not measured"] + [""] * (len(doc.tables[tno].columns) - 1)])

th = para_by_text("Rth,bond = [ TJ(∞) − TNTC(∞) ] / Pop")
add_block(th, [
 ("p", "Full thermometry was not run, since it needs a calibrated temperature stage to obtain "
       "the slope S. The pulse-length study below bounds the self-heating that would enter it. "
       "One red channel was swept three times with only the integration length changed, giving "
       "current-on times of 5, 20 and 80 ms per point."),
 ("fig", "fig8_self_heating.png"),
 ("cap", "Fig. 7. Extracted series resistance against current-on time, same channel throughout. "
         "Error bars are the standard error of the fit."),
 ("p", "Between 5 and 20 ms the extracted resistance is unchanged within the fit error. At 80 ms "
       "it falls by 1.49 Ω. The die dissipates 28.4 mW at 13.9 mA, and a junction-to-board "
       "thermal resistance of order 200 K/W gives about 6 K of rise; with dVF/dT near "
       "−2 mV/K that appears as roughly −1.2 Ω of apparent series resistance, which matches the "
       "measurement to within 25 %. All campaign data was taken at the 5 ms setting. The "
       "thermal time constant of the package therefore lies between 20 and 80 ms."),
])
fill(doc.tables[5], [["1–8", "S7 D1 R", "not calibrated", "28.4", "≈6 (inferred)",
                      "≈200 (order of magnitude)"]])

# ------------------------------------------------------------------- summary
rows_sum = []
for s in range(1, 9):
    v = [float(f["Rs"]) for f in bys.get(s, [])]
    vf = [float(f["vf10"]) for f in bys.get(s, [])]
    rows_sum.append([f"{s}", f"{100*ok[s]/tot[s]:.1f}", "n/a (chain fault)",
                     f"{mean(vf):.3f}" if vf else "—",
                     f"{mean(v):.2f}" if v else "—",
                     DOMINANT[s], "not measured"])
t8 = doc.tables[7]
t8.rows[0].cells[2].text = "Rbond (mΩ)"
t8.rows[0].cells[4].text = "Rs red (Ω)"
t8.rows[0].cells[5].text = "Dominant failure mode"
t8.rows[0].cells[6].text = "ρc / Rsh"
for cell in t8.rows[0].cells:
    for r in cell.paragraphs[0].runs: r.font.size = Pt(8)
fill(t8, rows_sum)

g = para_by_text("Table VIII collects the headline number")
add_block(g, [
 ("p", "Two results survive. Yield and failure mode separate the eight conditions at "
       "p = 2.2 × 10⁻⁴, and place conditions 5 and 7 ahead of the rest with no failures in "
       "12 channels each. Series resistance does not separate them, and the reason is measured: "
       "the fixture repeatability of 0.84 Ω exceeds the bond resistances by more than an order "
       "of magnitude."),
 ("p", "Recovering series resistance as a discriminator needs the sense connections moved to "
       "the Tier-1 probe pads so that the jumper contacts fall outside the measured loop. "
       "Nothing about the dice or the bonding prevents it; the round 2 wiring simply placed both "
       "taps on the wrong side of the connectors."),
])

# ------------------------------------------------------------------- table I
fill(doc.tables[0], [
 ["Bond / chain resistance", "Daisy chains DCL6, DCL12", "Not usable (layout fault)", "—"],
 ["Forward voltage, series resistance", "D1–D8, south header pins", "Arduino UNO R3, 6-branch "
  "binary current source, 14-bit oversampled ADC", "Pulsed DC, 63 levels, 0.5–15.9 mA, 5 ms on / 250 ms off"],
 ["Channel screening and failure mode", "D1–D8 probe pads", "Handheld DMM, diode test", "≈1 mA forward"],
 ["Sense resistor calibration", "R_EIS_LOAD, 100 Ω 0.1 %", "Handheld DMM, 600 Ω range, REL", "DC"],
 ["Specific contact resistivity", "TLM ladders", "Not measured", "—"],
 ["Sheet resistance", "van der Pauw cloverleaves", "Not measured", "—"]])

# relabel the subsection letters in document order
import string
h2 = [q for q in doc.paragraphs if q.style.name == "Heading 2"]
for idx, q in enumerate(h2):
    body = re.sub(r"^[A-Z]\.\s*", "", q.text.strip())
    for r_ in q.runs:
        r_.text = ""
    q.runs[0].text = f"{string.ascii_uppercase[idx]}.  {body}"

doc.save(os.path.join(OUT, "Electrical_Characterization_section.docx"))
print("docx written to", OUT)
print(f"red channels {len(red)}, yield chi2 p = 2.2e-4")
